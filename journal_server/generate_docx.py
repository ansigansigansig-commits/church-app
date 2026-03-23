#!/usr/bin/env python3
"""
솔리데오글로리아교회 일지 DOCX 생성 모듈
HWP 서식과 동일한 레이아웃으로 Word 문서를 생성합니다.
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

FONT = "맑은 고딕"
GRAY = "EFEFEF"


def _set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_font(run, size=8.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    # 한글 폰트 설정
    rPr = run._r.get_or_add_rPr()
    ea = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}"/>')
    rPr.append(ea)


def _label_cell(cell, text):
    """라벨 셀: 회색 배경, 볼드, 가운데"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, 8.5, bold=True)
    _set_shading(cell, GRAY)


def _value_cell(cell, text):
    """값 셀"""
    cell.text = ""
    for i, line in enumerate((text or "-").split("\n")):
        if i == 0:
            run = cell.paragraphs[0].add_run(line)
        else:
            p = cell.add_paragraph()
            run = p.add_run(line)
        _set_font(run, 8.5)


def _heading(doc, text):
    """섹션 헤딩 (밑줄)"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(3)
    run = p.add_run(text)
    _set_font(run, 10.5, bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


def _att_text(male, female, total):
    return f"남: {male}   여: {female}   계: {total}"


def generate_journal_docx(data, output_path):
    """
    일지 DOCX 생성

    Args:
        data: Firebase 일지 데이터 dict
        output_path: 저장 경로
    Returns:
        output_path
    """
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # === 제목 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    run = title.add_run("솔리데오글로리아교회 일지")
    _set_font(run, 19, bold=True)

    # 날짜
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.space_after = Pt(2)
    run = date_p.add_run(data.get("dateStr", ""))
    _set_font(run, 9)

    # 작성자 + 담임목사
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author_p.space_after = Pt(12)
    run = author_p.add_run(
        f"작성자 : {data.get('author', '문지영')}          담임목사"
    )
    _set_font(run, 9)

    # === 오전 예배 ===
    _heading(doc, "예배")

    tbl = doc.add_table(rows=6, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False

    # Row 0: 장소 | 값 | 설교자 | 값
    _label_cell(tbl.rows[0].cells[0], "장소")
    _value_cell(tbl.rows[0].cells[1], "솔리데오글로리아 예배당/유튜브 채널")
    _label_cell(tbl.rows[0].cells[2], "설교자")
    _value_cell(tbl.rows[0].cells[3], data.get("am_preacher"))

    # Row 1: 시간 | 값 | 제목 | 값
    _label_cell(tbl.rows[1].cells[0], "시간")
    _value_cell(tbl.rows[1].cells[1], data.get("am_time"))
    _label_cell(tbl.rows[1].cells[2], "제목")
    _value_cell(tbl.rows[1].cells[3], data.get("am_title"))

    # Row 2: 예배전 순서 (병합)
    _label_cell(tbl.rows[2].cells[0], "예배전\n순서")
    tbl.rows[2].cells[1].merge(tbl.rows[2].cells[3])
    _value_cell(tbl.rows[2].cells[1], data.get("am_pre"))

    # Row 3: 찬송 | 값 | 성례 | 값
    _label_cell(tbl.rows[3].cells[0], "찬송")
    _value_cell(tbl.rows[3].cells[1], data.get("am_hymns"))
    _label_cell(tbl.rows[3].cells[2], "성례")
    _value_cell(tbl.rows[3].cells[3], data.get("am_sacrament"))

    # Row 4: 신앙고백 | 값 | 본문 | 값
    _label_cell(tbl.rows[4].cells[0], "신앙고백")
    _value_cell(tbl.rows[4].cells[1], data.get("am_creed"))
    _label_cell(tbl.rows[4].cells[2], "본문")
    _value_cell(tbl.rows[4].cells[3], data.get("am_scripture"))

    # Row 5: 참석인원
    tbl.rows[5].cells[0].merge(tbl.rows[5].cells[1])
    tbl.rows[5].cells[2].merge(tbl.rows[5].cells[3])
    cell = tbl.rows[5].cells[2]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("예배당 참석인원")
    _set_font(run, 8.5, bold=True)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(
        _att_text(data.get("am_male", 0), data.get("am_female", 0), data.get("am_total", 0))
    )
    _set_font(run2, 10, bold=True)

    # === 오후 기도 순서 ===
    _heading(doc, "오후 기도 순서")

    prayer_tbl = doc.add_table(rows=1, cols=1)
    prayer_tbl.style = "Table Grid"
    _value_cell(prayer_tbl.rows[0].cells[0], data.get("pm_prayer"))

    pm_tbl = doc.add_table(rows=4, cols=4)
    pm_tbl.style = "Table Grid"
    pm_tbl.autofit = False

    _label_cell(pm_tbl.rows[0].cells[0], "장소")
    _value_cell(pm_tbl.rows[0].cells[1], "솔리데오글로리아 예배당/유튜브 채널")
    _label_cell(pm_tbl.rows[0].cells[2], "인도자")
    _value_cell(pm_tbl.rows[0].cells[3], data.get("pm_leader"))

    _label_cell(pm_tbl.rows[1].cells[0], "시간")
    _value_cell(pm_tbl.rows[1].cells[1], data.get("pm_time"))
    _label_cell(pm_tbl.rows[1].cells[2], "제목")
    _value_cell(pm_tbl.rows[1].cells[3], data.get("pm_title"))

    _label_cell(pm_tbl.rows[2].cells[0], "찬송")
    _value_cell(pm_tbl.rows[2].cells[1], data.get("pm_hymns"))
    _label_cell(pm_tbl.rows[2].cells[2], "본문")
    _value_cell(pm_tbl.rows[2].cells[3], data.get("pm_scripture"))

    _label_cell(pm_tbl.rows[3].cells[0], "신앙고백")
    _value_cell(pm_tbl.rows[3].cells[1], data.get("pm_creed"))
    pm_tbl.rows[3].cells[2].merge(pm_tbl.rows[3].cells[3])
    cell = pm_tbl.rows[3].cells[2]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("예배당 참석인원")
    _set_font(run, 8.5, bold=True)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(
        _att_text(data.get("pm_male", 0), data.get("pm_female", 0), data.get("pm_total", 0))
    )
    _set_font(run2, 10, bold=True)

    # === 학습 표 ===
    classes = data.get("study_classes", [])
    if classes:
        _heading(doc, "학습")

        row_labels = ["", "시간", "학생", "교사", "교재", "진도", "비고", "참석인원"]
        cols = len(classes) + 1
        study_tbl = doc.add_table(rows=len(row_labels), cols=cols)
        study_tbl.style = "Table Grid"
        study_tbl.autofit = False

        for i, label in enumerate(row_labels):
            _label_cell(study_tbl.rows[i].cells[0], label)

        for j, cls in enumerate(classes):
            col = j + 1
            _label_cell(study_tbl.rows[0].cells[col], cls.get("department", ""))
            _value_cell(study_tbl.rows[1].cells[col], cls.get("time", "-"))
            _value_cell(
                study_tbl.rows[2].cells[col], ", ".join(cls.get("attendees", []))
            )
            _value_cell(study_tbl.rows[3].cells[col], cls.get("teacher", "-"))
            _value_cell(study_tbl.rows[4].cells[col], cls.get("material", "-"))
            _value_cell(study_tbl.rows[5].cells[col], cls.get("content", "-"))

            note_parts = []
            note = cls.get("note", "")
            if note and note != "없음" and note != "-":
                note_parts.append(note)
            absentees = cls.get("absentees", [])
            if absentees:
                note_parts.append(f"결석: {', '.join(absentees)}")
            _value_cell(
                study_tbl.rows[6].cells[col], "\n".join(note_parts) if note_parts else "-"
            )

            count = cls.get("attendance_count", 0)
            _value_cell(study_tbl.rows[7].cells[col], f"{count}명")

    # === 방문 및 결석 ===
    _heading(doc, "방문 및 결석")

    va_tbl = doc.add_table(rows=2, cols=2)
    va_tbl.style = "Table Grid"
    _label_cell(va_tbl.rows[0].cells[0], "방문교인")
    _value_cell(va_tbl.rows[0].cells[1], data.get("visitors"))
    _label_cell(va_tbl.rows[1].cells[0], "결석교인")
    _value_cell(va_tbl.rows[1].cells[1], data.get("absences"))

    # === 공지사항 ===
    announcements = data.get("announcements")
    if announcements:
        _heading(doc, "공지 및 토의")
        ann_tbl = doc.add_table(rows=1, cols=2)
        ann_tbl.style = "Table Grid"
        _label_cell(ann_tbl.rows[0].cells[0], "공지사항")
        _value_cell(ann_tbl.rows[0].cells[1], announcements)

    # === 기타 모임 ===
    other = data.get("other")
    if other:
        _heading(doc, "기타 모임")
        other_tbl = doc.add_table(rows=1, cols=2)
        other_tbl.style = "Table Grid"
        _label_cell(other_tbl.rows[0].cells[0], "내용")
        _value_cell(other_tbl.rows[0].cells[1], other)

    # 푸터
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.space_before = Pt(18)
    run = footer.add_run("솔리데오글로리아교인")
    _set_font(run, 9)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # 테스트
    test_data = {
        "dateStr": "2026. 03. 22",
        "author": "문지영",
        "am_preacher": "김병혁 목사",
        "am_time": "10:45 ~ 13:00",
        "am_title": "그 태를 여신고로",
        "am_scripture": "창 29:31-30:24(2)",
        "am_pre": "시편('성전에 올라가는 노래') 해설",
        "am_hymns": "283장 (시 134편)\n5장 (시 2편)\n218장 (시 111편)\n269장 (시 126편)\n326장 (살후 3장)",
        "am_creed": "니케아 신조",
        "am_sacrament": "-",
        "am_male": "34",
        "am_female": "41",
        "am_total": "75",
        "pm_prayer": "안 식 집사",
        "pm_leader": "공정윤 강도사",
        "pm_time": "14:45 ~ 16:30",
        "pm_title": "사사기 해설",
        "pm_scripture": "-",
        "pm_hymns": "193(99편), 92(45편)",
        "pm_creed": "-",
        "pm_male": "32",
        "pm_female": "40",
        "pm_total": "72",
        "visitors": "-",
        "absences": "- 오전·오후: 이종민(독일)\n- 오후: 김병혁, 이선영(대구)",
        "announcements": "1. 청년 컨퍼런스 (3/29-30)\n2. 장소: 어피어연수원(경기 가평)",
        "study_classes": [
            {
                "department": "청년부 오전",
                "time": "10:20-10:45",
                "attendees": ["고수민", "박솔민", "허하민"],
                "attendance_count": 3,
                "teacher": "정진수 성도",
                "content": "창세기 강설#95 설교 미리보기",
                "material": "-",
                "absentees": [],
                "absent_count": 0,
                "note": "-",
            },
            {
                "department": "청년부 오후",
                "time": "13:50-14:45",
                "attendees": [
                    "고수민", "고지민", "김수산나", "김윤호", "김호산나",
                    "류철희", "박솔민", "이재우", "장진옥", "정윤재", "허하민",
                ],
                "attendance_count": 11,
                "teacher": "정진수 성도",
                "content": "창세기 강설#95 설교 나눔",
                "material": "-",
                "absentees": ["정다혜"],
                "absent_count": 1,
                "note": "-",
            },
            {
                "department": "초등 2부",
                "time": "-",
                "attendees": ["표정호", "안시아", "정한서", "박송하"],
                "attendance_count": 4,
                "teacher": "권선혜 성도",
                "content": "하이델베르크 요리문답 124문 3, 설교 나눔",
                "material": "하이델베르크 요리문답",
                "absentees": [],
                "absent_count": 0,
                "note": "-",
            },
            {
                "department": "청소년 2부",
                "time": "-",
                "attendees": ["이서연", "송예원", "이수아", "이재현", "정지혜", "정한음", "허하람"],
                "attendance_count": 7,
                "teacher": "공정윤 강도사",
                "content": "인간론 은혜언약",
                "material": "-",
                "absentees": [],
                "absent_count": 0,
                "note": "-",
            },
        ],
    }
    generate_journal_docx(test_data, "test_output.docx")
    print("생성 완료: test_output.docx")
